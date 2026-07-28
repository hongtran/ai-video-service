import io
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from fastapi import HTTPException, UploadFile

from app.api.router import request_screencast_video, request_video
from app.api.schemas import CreateVideoRequest
from app.config import Settings
from app.domain.models import JobStatus
from app.llm.client import StubScriptNormalizer
from app.storage.artifacts import LocalArtifactStore
from app.storage.jobs import InMemoryJobRepository

_GIF_BYTES = b"GIF89a" + b"\x00" * 64
_PNG_BYTES = b"\x89PNG\r\n\x1a\n" + b"\x00" * 64  # right size, wrong magic


def _build_docx(sections: list[dict]) -> bytes:
    """sections: [{heading, narration, gif}] -> a minimal .docx matching the
    real guide's shape (one Heading-1 section per step, one narration
    paragraph, one gif-naming caption paragraph)."""
    doc = Document()
    for section in sections:
        doc.add_paragraph(section["heading"], style="Heading 1")
        doc.add_paragraph(section["narration"])
        doc.add_paragraph(f"Illustration ({section['gif']})")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class RecordingQueue:
    def __init__(self) -> None:
        self.enqueued: list[str] = []

    async def enqueue(self, job_id: str) -> None:
        self.enqueued.append(job_id)


class ScreencastUploadTests(unittest.IsolatedAsyncioTestCase):
    async def asyncSetUp(self) -> None:
        self.settings = Settings(max_gif_bytes=1024, max_docx_bytes=1024 * 1024)
        self._tmp = tempfile.TemporaryDirectory()
        self.artifacts = LocalArtifactStore(Path(self._tmp.name))
        self.jobs = InMemoryJobRepository()
        self.queue = RecordingQueue()
        self.normalizer = StubScriptNormalizer()

    async def asyncTearDown(self) -> None:
        self._tmp.cleanup()

    def _request(self) -> SimpleNamespace:
        return SimpleNamespace(
            app=SimpleNamespace(
                state=SimpleNamespace(
                    settings=self.settings,
                    jobs=self.jobs,
                    artifacts=self.artifacts,
                    queue=self.queue,
                    guard=None,
                    normalizer=self.normalizer,
                )
            )
        )

    def _upload(self, data: bytes, name: str) -> UploadFile:
        return UploadFile(filename=name, file=io.BytesIO(data))

    def _docx_upload(self, sections: list[dict], name: str = "guide.docx") -> UploadFile:
        return self._upload(_build_docx(sections), name)

    def _gif_upload(self, name: str, data: bytes = _GIF_BYTES) -> UploadFile:
        return self._upload(data, name)

    async def test_happy_path_matches_by_filename_regardless_of_upload_order(self) -> None:
        sections = [
            {"heading": "Step 1", "narration": "Click Settings.", "gif": "a-open.gif"},
            {"heading": "Step 2", "narration": "Invite a teammate.", "gif": "b-invite.gif"},
            {"heading": "Step 3", "narration": "Confirm the invite.", "gif": "c-confirm.gif"},
        ]
        document = self._docx_upload(sections)
        # Uploaded SCRAMBLED — proves matching is by name, not upload order.
        files = [
            self._gif_upload("c-confirm.gif"),
            self._gif_upload("a-open.gif"),
            self._gif_upload("b-invite.gif"),
        ]

        response = await request_screencast_video(
            self._request(), document=document, files=files,
            subject="user-guide", orientation="horizontal", language="en",
        )

        self.assertEqual(response.status, JobStatus.PENDING)
        self.assertEqual(response.subject, "user-guide")
        self.assertEqual(self.queue.enqueued, [response.id])

        self.assertTrue(self.artifacts.exists(response.id, "source.docx"))
        self.assertTrue(self.artifacts.exists(response.id, "steps.json"))
        steps = self.artifacts.load_json(response.id, "steps.json")
        self.assertEqual([s["gif"] for s in steps], ["step-000.gif", "step-001.gif", "step-002.gif"])
        self.assertEqual(steps[0]["narration"], "Click Settings.")
        self.assertEqual(steps[1]["narration"], "Invite a teammate.")
        self.assertEqual(steps[2]["narration"], "Confirm the invite.")

        for i in range(3):
            self.assertTrue(self.artifacts.exists(response.id, f"source_gifs/step-{i:03d}.gif"))

        job = await self.jobs.get(response.id)
        self.assertIsNotNone(job)
        self.assertEqual(job.input_mode, "script")
        self.assertIn("Click Settings.", job.script)
        self.assertIn("Invite a teammate.", job.script)
        self.assertIn("Confirm the invite.", job.script)

    async def test_non_gif_magic_bytes_among_files_rejected_naming_the_file(self) -> None:
        sections = [{"heading": "Step 1", "narration": "Click Settings.", "gif": "step.gif"}]
        with self.assertRaises(HTTPException) as caught:
            await request_screencast_video(
                self._request(),
                document=self._docx_upload(sections),
                files=[self._gif_upload("step.gif", data=_PNG_BYTES)],
                subject="user-guide", orientation="horizontal", language="en",
            )

        self.assertEqual(caught.exception.status_code, 400)
        self.assertIn("step.gif", caught.exception.detail)
        self.assertEqual(self.queue.enqueued, [])
        self.assertEqual(await self.jobs.list(), [])

    async def test_oversized_gif_rejected_without_buffering_whole_file(self) -> None:
        sections = [{"heading": "Step 1", "narration": "Click Settings.", "gif": "step.gif"}]
        oversized = _GIF_BYTES + b"\x00" * (self.settings.max_gif_bytes * 2)
        with self.assertRaises(HTTPException) as caught:
            await request_screencast_video(
                self._request(),
                document=self._docx_upload(sections),
                files=[self._gif_upload("step.gif", data=oversized)],
                subject="user-guide", orientation="horizontal", language="en",
            )

        self.assertEqual(caught.exception.status_code, 413)
        self.assertEqual(self.queue.enqueued, [])

    async def test_non_docx_magic_bytes_rejected(self) -> None:
        with self.assertRaises(HTTPException) as caught:
            await request_screencast_video(
                self._request(),
                document=self._upload(_PNG_BYTES, "guide.docx"),
                files=[self._gif_upload("step.gif")],
                subject="user-guide", orientation="horizontal", language="en",
            )

        self.assertEqual(caught.exception.status_code, 400)
        self.assertEqual(self.queue.enqueued, [])

    async def test_oversized_docx_rejected(self) -> None:
        sections = [{"heading": "Step 1", "narration": "Click Settings.", "gif": "step.gif"}]
        small_settings = Settings(max_docx_bytes=16)
        request = self._request()
        request.app.state.settings = small_settings
        with self.assertRaises(HTTPException) as caught:
            await request_screencast_video(
                request,
                document=self._docx_upload(sections),
                files=[self._gif_upload("step.gif")],
                subject="user-guide", orientation="horizontal", language="en",
            )

        self.assertEqual(caught.exception.status_code, 413)

    async def test_doc_with_no_heading_rejected(self) -> None:
        doc = Document()
        doc.add_paragraph("Just plain text, no Heading style anywhere.")
        buf = BytesIO()
        doc.save(buf)

        with self.assertRaises(HTTPException) as caught:
            await request_screencast_video(
                self._request(),
                document=self._upload(buf.getvalue(), "guide.docx"),
                files=[self._gif_upload("step.gif")],
                subject="user-guide", orientation="horizontal", language="en",
            )

        self.assertEqual(caught.exception.status_code, 400)

    async def test_step_referencing_an_unuploaded_gif_rejected(self) -> None:
        sections = [
            {"heading": "Step 1", "narration": "Click Settings.", "gif": "expected.gif"},
        ]
        with self.assertRaises(HTTPException) as caught:
            await request_screencast_video(
                self._request(),
                document=self._docx_upload(sections),
                files=[self._gif_upload("wrong-name.gif")],
                subject="user-guide", orientation="horizontal", language="en",
            )

        self.assertEqual(caught.exception.status_code, 400)
        self.assertIn("expected.gif", caught.exception.detail)
        self.assertEqual(self.queue.enqueued, [])

    async def test_uploaded_gif_not_referenced_by_any_step_rejected(self) -> None:
        sections = [
            {"heading": "Step 1", "narration": "Click Settings.", "gif": "step-1.gif"},
        ]
        with self.assertRaises(HTTPException) as caught:
            await request_screencast_video(
                self._request(),
                document=self._docx_upload(sections),
                files=[self._gif_upload("step-1.gif"), self._gif_upload("extra.gif")],
                subject="user-guide", orientation="horizontal", language="en",
            )

        self.assertEqual(caught.exception.status_code, 400)
        self.assertIn("extra.gif", caught.exception.detail)

    async def test_no_gif_files_uploaded_rejected(self) -> None:
        sections = [{"heading": "Step 1", "narration": "Click Settings.", "gif": "step.gif"}]
        with self.assertRaises(HTTPException) as caught:
            await request_screencast_video(
                self._request(),
                document=self._docx_upload(sections),
                files=[],
                subject="user-guide", orientation="horizontal", language="en",
            )

        self.assertEqual(caught.exception.status_code, 400)


class UserGuideRejectedOnPlainVideosRouteTests(unittest.IsolatedAsyncioTestCase):
    """user-guide cannot work without a docx + GIFs, so the JSON-body /videos
    route (which has no file upload) must refuse it rather than silently
    trying."""

    async def test_user_guide_subject_rejected_on_post_videos(self) -> None:
        settings = Settings()
        jobs = InMemoryJobRepository()
        queue = RecordingQueue()

        request = SimpleNamespace(
            app=SimpleNamespace(
                state=SimpleNamespace(
                    settings=settings,
                    jobs=jobs,
                    artifacts=None,
                    queue=queue,
                    guard=None,
                    normalizer=StubScriptNormalizer(),
                )
            )
        )

        with self.assertRaises(HTTPException) as caught:
            await request_video(
                CreateVideoRequest(
                    input_mode="script",
                    script="Click Settings.",
                    subject="user-guide",
                    orientation="horizontal",
                ),
                request,
            )

        self.assertEqual(caught.exception.status_code, 400)
        self.assertEqual(queue.enqueued, [])


if __name__ == "__main__":
    unittest.main()
