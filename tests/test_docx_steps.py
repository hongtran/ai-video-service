import unittest
from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from app.documents.docx import DocxExtractError, extract_steps


def _set_outline_level(paragraph, level: int) -> None:
    """Apply a w:outlineLvl directly to a paragraph, the way Word does when an author
    sets an outline level without applying a named Heading style."""
    pPr = paragraph._p.get_or_add_pPr()
    element = pPr.makeelement(qn("w:outlineLvl"), {qn("w:val"): str(level)})
    pPr.append(element)


def _save(doc) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx(sections: list[dict]) -> bytes:
    """A minimal in-memory .docx: each section is
    {heading, narration_lines: [...], captions: [...]}, written as a
    Heading-1 paragraph followed by plain paragraphs — python-docx both
    writes and reads, so no fixture file is needed."""
    doc = Document()
    for section in sections:
        doc.add_paragraph(section["heading"], style="Heading 1")
        for line in section.get("narration_lines", []):
            doc.add_paragraph(line)
        for caption in section.get("captions", []):
            doc.add_paragraph(caption)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class ExtractStepsTests(unittest.TestCase):
    def test_two_steps_extracted_in_order(self) -> None:
        data = _build_docx([
            {
                "heading": "Step 1: Open settings",
                "narration_lines": ["Click the gear icon in the top right."],
                "captions": ["Illustration Step 1 (step-1-open-settings.gif)"],
            },
            {
                "heading": "Step 2: Invite a teammate",
                "narration_lines": ["Enter their email and click Send."],
                "captions": ["Illustration Step 2 (step-2-invite.gif)"],
            },
        ])

        steps = extract_steps(data)

        self.assertEqual(len(steps), 2)
        self.assertEqual(steps[0].gif_filename, "step-1-open-settings.gif")
        self.assertEqual(steps[0].narration, "Click the gear icon in the top right.")
        self.assertEqual(steps[1].gif_filename, "step-2-invite.gif")
        self.assertEqual(steps[1].narration, "Enter their email and click Send.")

    def test_heading_and_caption_text_never_appear_in_narration(self) -> None:
        data = _build_docx([{
            "heading": "Step 1: Open settings",
            "narration_lines": ["Click the gear icon."],
            "captions": ["Illustration Step 1 (open-settings.gif)"],
        }])

        narration = extract_steps(data)[0].narration
        self.assertNotIn("Step 1", narration)
        self.assertNotIn("Illustration", narration)
        self.assertNotIn(".gif", narration)

    def test_multiple_narration_paragraphs_are_joined(self) -> None:
        data = _build_docx([{
            "heading": "Step 1",
            "narration_lines": ["First sentence.", "Second sentence."],
            "captions": ["(step-1.gif)"],
        }])

        narration = extract_steps(data)[0].narration
        self.assertIn("First sentence.", narration)
        self.assertIn("Second sentence.", narration)

    def test_front_matter_before_first_heading_is_ignored(self) -> None:
        doc = Document()
        doc.add_paragraph("A Guide", style="Title")
        doc.add_paragraph("Some subtitle text.")
        doc.add_paragraph("Step 1", style="Heading 1")
        doc.add_paragraph("Narration.")
        doc.add_paragraph("(step-1.gif)")
        buf = BytesIO()
        doc.save(buf)

        steps = extract_steps(buf.getvalue())
        self.assertEqual(len(steps), 1)
        self.assertEqual(steps[0].narration, "Narration.")

    def test_no_heading_anywhere_raises(self) -> None:
        doc = Document()
        doc.add_paragraph("Just a plain paragraph, no heading style at all.")
        buf = BytesIO()
        doc.save(buf)

        with self.assertRaises(DocxExtractError):
            extract_steps(buf.getvalue())

    def test_step_with_no_gif_reference_raises(self) -> None:
        data = _build_docx([{
            "heading": "Step 1",
            "narration_lines": ["Click the button."],
            "captions": [],
        }])

        with self.assertRaises(DocxExtractError):
            extract_steps(data)

    def test_step_with_two_distinct_gif_references_raises(self) -> None:
        data = _build_docx([{
            "heading": "Step 1",
            "narration_lines": ["Click the button."],
            "captions": ["(step-1a.gif)", "(step-1b.gif)"],
        }])

        with self.assertRaises(DocxExtractError):
            extract_steps(data)

    def test_step_with_repeated_identical_gif_reference_is_fine(self) -> None:
        # Same filename mentioned twice (e.g. a before/after caption) is not
        # ambiguous — only genuinely DIFFERENT filenames are.
        data = _build_docx([{
            "heading": "Step 1",
            "narration_lines": ["Click the button."],
            "captions": ["(step-1.gif)", "See (step-1.gif) again"],
        }])

        steps = extract_steps(data)
        self.assertEqual(steps[0].gif_filename, "step-1.gif")

    def test_step_with_only_heading_and_caption_raises(self) -> None:
        data = _build_docx([{
            "heading": "Step 1",
            "narration_lines": [],
            "captions": ["(step-1.gif)"],
        }])

        with self.assertRaises(DocxExtractError):
            extract_steps(data)

    def test_not_a_docx_raises(self) -> None:
        with self.assertRaises(DocxExtractError):
            extract_steps(b"not a docx file at all")

    def test_five_step_shape_matching_real_sample(self) -> None:
        # Mirrors the actual structure of Huong-dan-su-dung-Co-cau-to-chuc.docx:
        # 5 Heading-1 sections, one narration paragraph and one gif caption each.
        sections = [
            {
                "heading": f"Bước {i}: Step title {i}",
                "narration_lines": [f"Narration text for step {i}."],
                "captions": [f"Ảnh minh họa Bước {i} (buoc-{i}-demo.gif)"],
            }
            for i in range(1, 6)
        ]
        data = _build_docx(sections)

        steps = extract_steps(data)

        self.assertEqual(len(steps), 5)
        for i, step in enumerate(steps, start=1):
            self.assertEqual(step.gif_filename, f"buoc-{i}-demo.gif")
            self.assertEqual(step.narration, f"Narration text for step {i}.")


class HeadingDetectionTests(unittest.TestCase):
    """Documents whose step structure is real but not expressed as a literal
    English "Heading 1" style — each of these used to raise "no step headings"."""

    def test_custom_style_based_on_heading_1_splits_steps(self) -> None:
        doc = Document()
        style = doc.styles.add_style("Bước", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Heading 1"]
        for i in (1, 2):
            doc.add_paragraph(f"Bước {i}", style=style)
            doc.add_paragraph(f"Narration {i}.")
            doc.add_paragraph(f"(step-{i}.gif)")

        steps = extract_steps(_save(doc))

        self.assertEqual([s.gif_filename for s in steps], ["step-1.gif", "step-2.gif"])
        self.assertEqual(steps[0].narration, "Narration 1.")

    def test_localized_style_name_with_english_style_id_splits_steps(self) -> None:
        # Word keeps the built-in styleId ("Heading1") in the XML while the UI shows a
        # localized name, and some editors write that localized name into styles.xml.
        doc = Document()
        heading = doc.styles["Heading 1"]
        heading.name = "Đề mục 1"
        doc.add_paragraph("Bước 1", style=heading)
        doc.add_paragraph("Narration.")
        doc.add_paragraph("(step-1.gif)")

        steps = extract_steps(_save(doc))

        self.assertEqual(len(steps), 1)
        self.assertEqual(steps[0].narration, "Narration.")

    def test_paragraph_level_outline_level_splits_steps(self) -> None:
        doc = Document()
        title = doc.add_paragraph("Bước 1")  # Normal style, outline level only
        _set_outline_level(title, 0)
        doc.add_paragraph("Narration.")
        doc.add_paragraph("(step-1.gif)")

        steps = extract_steps(_save(doc))

        self.assertEqual(len(steps), 1)
        self.assertEqual(steps[0].gif_filename, "step-1.gif")

    def test_outline_level_9_is_body_text_not_a_heading(self) -> None:
        # Word writes outlineLvl 9 to mean "body text"; it must not start a step.
        doc = Document()
        doc.add_paragraph("Bước 1", style="Heading 1")
        body = doc.add_paragraph("Narration.")
        _set_outline_level(body, 9)
        doc.add_paragraph("(step-1.gif)")

        steps = extract_steps(_save(doc))

        self.assertEqual(len(steps), 1)
        self.assertEqual(steps[0].narration, "Narration.")

    def test_sub_headings_stay_inside_their_parent_step(self) -> None:
        doc = Document()
        doc.add_paragraph("Bước 1", style="Heading 1")
        doc.add_paragraph("Chi tiết", style="Heading 2")
        doc.add_paragraph("Narration one.")
        doc.add_paragraph("(step-1.gif)")
        doc.add_paragraph("Bước 2", style="Heading 1")
        doc.add_paragraph("Narration two.")
        doc.add_paragraph("(step-2.gif)")

        steps = extract_steps(_save(doc))

        self.assertEqual(len(steps), 2)
        self.assertIn("Chi tiết", steps[0].narration)
        self.assertIn("Narration one.", steps[0].narration)
        self.assertEqual(steps[1].narration, "Narration two.")

    def test_toc_heading_does_not_start_a_step(self) -> None:
        # "TOC Heading" is based on Heading 1, so the base-style walk would otherwise
        # treat a table-of-contents header as a step with no .gif reference.
        doc = Document()
        toc = doc.styles["TOC Heading"]  # ships with the default template
        self.assertEqual(toc.base_style.name, "Heading 1")
        doc.add_paragraph("Mục lục", style=toc)
        doc.add_paragraph("Bước 1", style="Heading 1")
        doc.add_paragraph("Narration.")
        doc.add_paragraph("(step-1.gif)")

        steps = extract_steps(_save(doc))

        self.assertEqual(len(steps), 1)
        self.assertEqual(steps[0].narration, "Narration.")

    def test_headings_and_narration_inside_a_table_are_found(self) -> None:
        # A guide laid out as a table is invisible to doc.paragraphs.
        doc = Document()
        table = doc.add_table(rows=2, cols=1)
        for row, i in zip(table.rows, (1, 2)):
            cell = row.cells[0]
            cell.paragraphs[0].text = f"Bước {i}"
            cell.paragraphs[0].style = doc.styles["Heading 1"]
            cell.add_paragraph(f"Narration {i}.")
            cell.add_paragraph(f"(step-{i}.gif)")

        steps = extract_steps(_save(doc))

        self.assertEqual([s.gif_filename for s in steps], ["step-1.gif", "step-2.gif"])
        self.assertEqual(steps[1].narration, "Narration 2.")


class UnstyledHeadingFallbackTests(unittest.TestCase):
    """Documents where step titles were made bold and large by hand instead of being
    given a Heading style — structurally invisible, rescued by their text."""

    def test_numbered_step_titles_in_plain_paragraphs_split_steps(self) -> None:
        doc = Document()
        doc.add_paragraph("HƯỚNG DẪN SỬ DỤNG", style="Title")
        for i in (1, 2):
            doc.add_paragraph(f"Bước {i}: Làm việc gì đó")
            doc.add_paragraph(f"Narration {i}.")
            doc.add_paragraph(f"Ảnh minh họa Bước {i} (buoc-{i}.gif)")

        steps = extract_steps(_save(doc))

        self.assertEqual([s.gif_filename for s in steps], ["buoc-1.gif", "buoc-2.gif"])
        self.assertEqual(steps[0].narration, "Narration 1.")
        self.assertNotIn("Bước 1", steps[0].narration)

    def test_bare_numeric_titles_split_steps(self) -> None:
        doc = Document()
        for i in (1, 2):
            doc.add_paragraph(f"{i}. Open the settings panel")
            doc.add_paragraph(f"Narration {i}.")
            doc.add_paragraph(f"(step-{i}.gif)")

        steps = extract_steps(_save(doc))

        self.assertEqual(len(steps), 2)
        self.assertEqual(steps[1].narration, "Narration 2.")

    def test_fallback_never_applies_to_a_properly_headed_document(self) -> None:
        # A narration line beginning "1. " must not split a document that already has
        # real headings — the fallback only runs when nothing structural was found.
        data = _build_docx([{
            "heading": "Bước 1",
            "narration_lines": ["1. Nhấn nút Thêm.", "2. Điền biểu mẫu."],
            "captions": ["(step-1.gif)"],
        }])

        steps = extract_steps(data)

        self.assertEqual(len(steps), 1)
        self.assertIn("1. Nhấn nút Thêm.", steps[0].narration)
        self.assertIn("2. Điền biểu mẫu.", steps[0].narration)

    def test_long_numbered_narration_line_is_not_a_fallback_title(self) -> None:
        # In fallback mode a long line starting with a number is prose, not a title.
        long_line = "1. " + ("chi tiết rất dài " * 12)
        doc = Document()
        doc.add_paragraph("Bước 1: Mở cài đặt")
        doc.add_paragraph(long_line)
        doc.add_paragraph("(step-1.gif)")

        steps = extract_steps(_save(doc))

        self.assertEqual(len(steps), 1)
        self.assertIn(long_line.strip(), steps[0].narration)

    def test_error_message_lists_the_styles_actually_found(self) -> None:
        doc = Document()
        doc.add_paragraph("A Guide", style="Title")
        doc.add_paragraph("Just prose with no step structure whatsoever.")

        with self.assertRaises(DocxExtractError) as ctx:
            extract_steps(_save(doc))

        message = str(ctx.exception)
        self.assertIn("Title", message)
        self.assertIn("Normal", message)


if __name__ == "__main__":
    unittest.main()
